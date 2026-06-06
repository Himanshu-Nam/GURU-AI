import streamlit as st
from dotenv import load_dotenv
import os
import tempfile

# LangChain imports
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.documents import Document

# BM25
from rank_bm25 import BM25Okapi

# DOCX (IMPORTANT FIX)
from docx import Document as DocxDocument

# Load env
load_dotenv()

# -----------------------------
# STREAMLIT CONFIG
# -----------------------------
st.set_page_config(page_title="GURU AI", page_icon="🤖", layout="wide")

st.title("🤖 GURU AI - Advanced RAG Chatbot")

# -----------------------------
# SESSION STATE
# -----------------------------
if "messages" not in st.session_state:
    st.session_state.messages = []

if "bm25" not in st.session_state:
    st.session_state.bm25 = None

if "docs" not in st.session_state:
    st.session_state.docs = None

if "vectordb" not in st.session_state:
    st.session_state.vectordb = None


# -----------------------------
# GUARDRAILS
# -----------------------------
def guardrail_check(query):
    blocked_words = [
        "ignore previous instructions",
        "system prompt",
        "jailbreak",
        "reveal prompt",
        "act as system"
    ]
    return not any(word in query.lower() for word in blocked_words)


# -----------------------------
# FILE LOADER
# -----------------------------
def load_document(file_path, file_type):
    docs = []

    if file_type == "pdf":
        loader = PyPDFLoader(file_path)
        docs = loader.load()

    elif file_type == "txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        docs = [Document(page_content=text, metadata={"source": file_path})]

    elif file_type == "docx":
        doc = DocxDocument(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        docs = [Document(page_content=text, metadata={"source": file_path})]

    return docs


# -----------------------------
# BM25 BUILD
# -----------------------------
def build_bm25(docs):
    texts = [d.page_content for d in docs]
    tokenized = [t.lower().split() for t in texts]
    return BM25Okapi(tokenized)


# -----------------------------
# HYBRID RETRIEVAL (FIXED)
# -----------------------------
def hybrid_retrieve(query):
    retriever = st.session_state.vectordb.as_retriever(
        search_type="mmr",
        search_kwargs={"k": 4, "fetch_k": 10}
    )

    # FIX: new LangChain API
    vector_docs = retriever.invoke(query)

    # BM25
    bm25 = st.session_state.bm25
    scores = bm25.get_scores(query.lower().split())

    top_idx = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)[:4]
    bm25_docs = [st.session_state.docs[i] for i in top_idx]

    # merge + deduplicate
    seen = set()
    final_docs = []

    for d in vector_docs + bm25_docs:
        if d.page_content not in seen:
            final_docs.append(d)
            seen.add(d.page_content)

    return final_docs[:6]


# -----------------------------
# FORMAT CONTEXT
# -----------------------------
def format_docs(docs):
    return "\n\n".join(d.page_content for d in docs)


# -----------------------------
# LLM + PROMPT
# -----------------------------
model = ChatOpenAI(model="gpt-4o-mini", temperature=0.3)

prompt = ChatPromptTemplate.from_messages([
    ("system", """
You are a smart AI assistant.

Rules:
- Use ONLY the provided context
- Use chat history if needed
- If answer not found, say "I don't know based on the document"
"""),
    ("human", """
Chat History:
{history}

Context:
{context}

Question:
{question}
""")
])

rag_chain = prompt | model | StrOutputParser()


# -----------------------------
# SIDEBAR UPLOAD
# -----------------------------
with st.sidebar:
    st.header("📄 Upload Document")

    uploaded_file = st.file_uploader(
        "Upload PDF / TXT / DOCX",
        type=["pdf", "txt", "docx"]
    )

    if uploaded_file:

        file_type = uploaded_file.name.split(".")[-1]

        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_type}") as tmp:
            tmp.write(uploaded_file.read())
            file_path = tmp.name

        with st.spinner("Processing document..."):

            docs = load_document(file_path, file_type)

            splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200
            )

            split_docs = splitter.split_documents(docs)

            embedding = OpenAIEmbeddings()

            vectordb = Chroma.from_documents(
                documents=split_docs,
                embedding=embedding
            )

            bm25 = build_bm25(split_docs)

            st.session_state.vectordb = vectordb
            st.session_state.bm25 = bm25
            st.session_state.docs = split_docs

        st.success("Document processed successfully!")


# -----------------------------
# CHAT HISTORY
# -----------------------------
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.write(msg["content"])


# -----------------------------
# INPUT
# -----------------------------
question = st.chat_input("Ask anything from your document...")

if question:

    if not guardrail_check(question):
        st.warning("🚫 Unsafe query detected")
        st.stop()

    st.session_state.messages.append({
        "role": "user",
        "content": question
    })

    with st.chat_message("user"):
        st.write(question)

    history = "\n".join(
        [f"{m['role']}: {m['content']}" for m in st.session_state.messages[-6:]]
    )

    with st.spinner("Thinking..."):

        docs = hybrid_retrieve(question)
        context = format_docs(docs)

        response = rag_chain.invoke({
            "question": question,
            "context": context,
            "history": history
        })

    st.session_state.messages.append({
        "role": "assistant",
        "content": response
    })

    with st.chat_message("assistant"):
        st.write(response)