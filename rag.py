from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
import docx
import streamlit as st
import os

def load_document(file_path, file_type):

    docs = []

    if file_type == "pdf":
        loader = PyPDFLoader(file_path)
        docs = loader.load()

    elif file_type == "txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        docs = [Document(page_content=text, metadata={"source": file_path})]

    elif file_type == "docx":
        doc = docx.Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        docs = [Document(page_content=text, metadata={"source": file_path})]

    return docs

def split_documents(docs):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=150
    )

    return splitter.split_documents(docs)

def create_vector_db(chunks, persist_dir="chroma_db"):

    embedding_model = OpenAIEmbeddings()

    vector_db = Chroma.from_documents(
        documents=chunks,
        embedding=embedding_model,
        persist_directory=persist_dir
    )

    return vector_db
file_path,file_type=None,None
docs = load_document(file_path, file_type)

chunks = split_documents(docs)

vector_db = create_vector_db(chunks)

st.session_state.vectordb = vector_db
st.session_state.docs = chunks